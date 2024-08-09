from django.shortcuts import render, get_object_or_404
import datetime
from django.shortcuts import render, redirect
from datetime import datetime as dt
import xlwt
from django.core.files import File
import shutil, os
from pandas import DataFrame
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
import threading
import traceback
import random
from django.db.models import Count
from django.db import transaction
from django.core.files.storage import FileSystemStorage
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
import decimal
from filter_prospects.filters import filter_query
from filter_prospects.forms import ProspectForm
from filter_prospects.models import *
from .models import *
from django.db.models import Q
import pandas as pd
from marketing_machine.models import MarketingSequence, MarketingPlan, MarketingCampaign
from user.decorators import user_has_list_management_Permission, user_has_skip_trace_Permission, user_has_single_skip_trace_Balance
from prospectx_new import settings
from prospectx_new.settings import BASE_DIR
import stripe
from payments.views import is_subscribed
import traceback
import time
import base64
import os
import requests
import json
from docx import Document
from docx.shared import Inches
from django.contrib import messages
from user.models import *
import openpyxl
import pandas as pd
import csv
from user.models import Revenue


@user_has_skip_trace_Permission
def skip_trace_files_list(request, pk):
    skip_trace_data_list = []
    price_data = UserProfile.objects.filter(user=request.user)
    try:
        PrepaidBalance.objects.create(user=request.user, amount=0.00)
    except:
        pass
    obj = PrepaidBalance.objects.filter(user=request.user)
    if pk == 'list':
        if request.method == 'POST':
            name_search = request.POST.get("file_name_search")
            status_search = request.POST.getlist("status_search")
            total_records_search = request.POST.get("total_records_search")
            total_hits_search = request.POST.get("total_hits_search")
            created_date = request.POST.get("datepicker1")
            if created_date:
                # created_date = created_date.replace('/', '-')
                created_date = (datetime.strptime(created_date, "%Y-%m-%d")).date()
            try:
                skip_trace = SkipTraceFile.objects.filter(user=request.user).select_related('user').order_by('-id')

                if name_search:
                    skip_trace = [skip for skip in skip_trace if skip.file_name and skip.file_name.find(name_search) >= 0]
                if status_search:
                    skip_trace = [skip for skip in skip_trace if skip.status in status_search]
                if total_records_search:
                    skip_trace = [skip for skip in skip_trace if skip.total_records == str(total_records_search)]
                if total_hits_search:
                    skip_trace = [skip for skip in skip_trace if skip.total_hits == str(total_hits_search)]
                if created_date:
                    skip_trace = [skip for skip in skip_trace if skip.created_at == created_date]

                for skip in skip_trace:
                    if skip.existing_file:
                        export_link = 'media' + skip.file_name
                    elif skip.status == 'Uploaded':
                        export_link = 'media/exported_skiptrace/' + skip.file_name,
                    else:
                        export_link = 'javascript:void(0)'
                    data = {
                        "id": skip.id,
                        "file_name": skip.file_name,
                        "list_link": 'media' + skip.file_name if skip.existing_file else 'media/skiptrace_files/' + skip.file_name,
                        "exported_link": export_link,
                        "status": skip.status,
                        "total_records": skip.total_records,
                        "existing_matches": skip.existing_matches,
                        "existing_match_savings": skip.existing_match_savings,
                        "total_hits": skip.total_hits,
                        "hits_percentage": skip.hits_percentage,
                        "total_cost": skip.total_cost,
                        "is_process_complete": skip.is_process_complete,
                        "date": skip.created_at,
                    }
                    skip_trace_data_list.append(data)
                return render(request, 'skiptrace/skiptrace_list.html', {"data": skip_trace_data_list,
                                                                         "user": request.user, "obj": obj[0]})
            except:
                print(traceback.format_exc())
                return redirect('/skiptrace/skip_trace/list/')

        skip_trace = SkipTraceFile.objects.filter(user=request.user).order_by('-id')
        for skip in skip_trace:
            if skip.existing_file:
                export_link = 'media'+skip.file_name
            elif skip.status == 'Uploaded':
                export_link = 'media/exported_skiptrace/'+skip.file_name,
            else:
                export_link = 'javascript:void(0)'
            data = {
                "id": skip.id,
                "file_name": skip.file_name,
                "list_link": 'media'+skip.file_name if skip.existing_file else 'media/skiptrace_files/'+skip.file_name,
                "exported_link": export_link,
                "status": skip.status,
                "total_records": skip.total_records,
                "existing_matches": skip.existing_matches,
                "existing_match_savings": skip.existing_match_savings,
                "total_hits": skip.total_hits,
                "hits_percentage": skip.hits_percentage,
                "total_cost": skip.total_cost,
                "is_process_complete": skip.is_process_complete,
                "date": skip.created_at,
            }
            skip_trace_data_list.append(data)
        return render(request, 'skiptrace/skiptrace_list.html', {"data": skip_trace_data_list,
                                                                 "user": request.user, "obj": obj[0]})
    else:
        if request.method == 'POST':
            name_search = request.POST.get("file_name_search")
            # email_search = request.POST.get("email_search")
            # phone_number_id = request.POST.get("phone-number-id")
            # address_search = request.POST.get("address_search")
            property_address_search = request.POST.get("property_address_search")
            property_city_search = request.POST.get("property_city_search")
            property_state_search = request.POST.get("property_state_search")
            property_zip_search = request.POST.get("property_zip_search")
            try:
                skip_trace = SingleSkipTrace.objects.filter(user=request.user).select_related('user').order_by('-id')

                if name_search:
                    skip_trace = [skip for skip in skip_trace if skip.full_name and skip.full_name.find(name_search) >= 0]
                # if email_search:
                #     skip_trace = [skip for skip in skip_trace if skip.status in status_search]
                # if phone_number_id:
                #     skip_trace = [skip for skip in skip_trace if skip.total_records == str(total_records_search)]
                # if address_search:
                #     skip_trace = [skip for skip in skip_trace if skip.total_hits == str(total_hits_search)]
                if property_address_search:
                    skip_trace = [skip for skip in skip_trace if skip.property_address == property_address_search]
                if property_city_search:
                    skip_trace = [skip for skip in skip_trace if skip.property_city == property_city_search]
                if property_state_search:
                    skip_trace = [skip for skip in skip_trace if (skip.property_state == property_state_search or skip.property_state == property_state_search.upper())]
                if property_zip_search:
                    skip_trace = [skip for skip in skip_trace if skip.property_zip == property_zip_search]

                for skips in skip_trace:
                    data = {
                        "id": skips.id,
                        "file_name": skips.file_name,
                        "list_link": 'media/single_skiptrace/' + skips.file_name,
                        "full_name": skips.full_name,
                        "property_address": skips.property_address,
                        "property_city": skips.property_city,
                        "property_state": skips.property_state,
                        "property_zip": skips.property_zip,
                        "email": EmailTraced.objects.filter(skiptrace=skips.id)[0] if EmailTraced.objects.filter(skiptrace=skips.id).exists() else " ",
                        "phone": PhoneTraced.objects.filter(skiptrace=skips.id)[0] if PhoneTraced.objects.filter(skiptrace=skips.id).exists() else " ",
                        "address": AddressTraced.objects.filter(skiptrace=skips.id)[0] if AddressTraced.objects.filter(skiptrace=skips.id).exists() else " "
                    }
                    skip_trace_data_list.append(data)
                return render(request, 'skiptrace/single_skiptrace_list.html', {"data": skip_trace_data_list,
                                                                                "user": request.user,
                                                                                "obj": obj[0],
                                                                                "price_data": price_data[
                                                                                    0].skiptrace_price})
            except:
                print(traceback.format_exc())
                return redirect('/skiptrace/skip_trace/single/')
        skip_trace = SingleSkipTrace.objects.filter(user=request.user).order_by('-id')
        for skips in skip_trace:
            data = {
                "id": skips.id,
                "file_name": skips.file_name,
                "list_link": 'media/single_skiptrace/' + skips.file_name,
                "full_name": skips.full_name,
                "property_address": skips.property_address,
                "property_city": skips.property_city,
                "property_state": skips.property_state,
                "property_zip": skips.property_zip,
                "email": [s.email for s in EmailTraced.objects.filter(skiptrace=skips.id)],
                "phone": [p.phone for p in PhoneTraced.objects.filter(skiptrace=skips.id)],
                "address": [a.address for a in AddressTraced.objects.filter(skiptrace=skips.id)]
            }
            skip_trace_data_list.append(data)
        return render(request, 'skiptrace/single_skiptrace_list.html', {"data": skip_trace_data_list,
                                                                        "user": request.user,
                                                                        "obj": obj[0],
                                                                        "price_data": price_data[0].skiptrace_price})


def upload_skiptrace_file(request):
    return render(request, 'skiptrace/upload_skiptrace.html')


def success_purchase_small_plan(request):
    if request.method == 'POST':
        balance = request.POST.get('bal_id')
        where_to = request.POST.get('where_to')
        token = request.POST['stripeToken']
        try:
            with transaction.atomic():
                obj = UserStripeDetail.objects.get(user=request.user)
                source = stripe.Customer.create_source(
                    obj.customer_id,
                    source=token,
                )

                charge = stripe.Charge.create(
                    amount=balance,
                    currency='usd',
                    customer=obj.customer_id,
                    source=source.id,
                )
                stripe.InvoiceItem.create(
                    customer=obj.customer_id,
                    amount=balance,
                    currency='usd',
                    description='One-time Skip Trace Charge'
                )
                sk = stripe.Invoice.create(
                    customer=obj.customer_id,
                    auto_advance=True  # auto-finalize this draft after ~1 hour
                )
                stripe.Invoice.finalize_invoice(sk.id)

                # if PrepaidBalance.objects.filter(user=request.user).exists():
                prepaid_balance = PrepaidBalance.objects.get(user=request.user)
                prepaid_balance.amount = prepaid_balance.amount + (decimal.Decimal(balance)/100)
                prepaid_balance.save()
                s_amount = decimal.Decimal(balance)/100
                Revenue.objects.create(amount=s_amount)
        except:
            print(traceback.format_exc())
            # else:
            #     PrepaidBalance.objects.create(user=request.user, amount=(decimal.Decimal(balance)/100))
        if where_to == 'bulk':
            return redirect('skip_trace/list/')
        else:
            return redirect('skip_trace/single/')
    print("in get method")
    return render(request, 'skiptrace/single_skiptrace_list.html', {"data": "",
                                                                    "user": request.user,
                                                                    "obj": ""})


class GetSingleSkipTrace(APIView):
    # permission_classes = (IsAuthenticated,)
    def post(self, request):
        first_name = request.data.get('first_name')
        last_name = request.data.get('last_name')
        full_address = request.data.get('address')
        property_city = request.data.get('city')
        property_state = request.data.get('state')
        property_zip = request.data.get('zip')
        mailing_address = request.data.get('mailing_address')
        mailing_city = request.data.get('mailing_city')
        mailing_state = request.data.get('mailing_state')
        mailing_zip = request.data.get('mailing_zip')
        user = request.data.get('user')
        print("data is: ", first_name, last_name, full_address, property_city, property_state, property_zip, user, mailing_address, mailing_city, mailing_state, mailing_zip)
        if full_address:
            single_obj = SingleSkipTrace()
            single_obj.user_id = user
            single_obj.first_name = first_name
            single_obj.last_name = last_name
            single_obj.property_address = full_address

            single_obj.property_city = property_city
            single_obj.property_state = property_state
            single_obj.property_zip = property_zip
            single_obj.mailing_address = mailing_address
            single_obj.mailing_city = mailing_city
            single_obj.mailing_state = mailing_state
            single_obj.mailing_zip = mailing_zip
            try:
                url1 = "https://login-api.idicore.com/apiclient"
                payload = "{\"glba\":\"otheruse\",\"dppa\":\"none\"}"
                d = base64.b64encode(
                    b'api-client@acesolutions:ELw%7UYUbfkb5V!naYXhksX2Tu3LYvA%dqvA5YkevtXU%4x59GLnvhnJ5kqiL4Bi')
                print("d is: ", d)
                abc = "Basic " + str(d.decode("utf-8"))
                print("abc is: ", abc)
                headers = {
                    'authorization': abc,
                    'content-type': "application/json"
                }
                response = requests.request("POST", url1, data=payload, headers=headers)
                my_token = response.text  # token created for 15 minutes
                print("skkiptrace token: ", my_token)
                url2 = "https://api.idicore.com/search"
                headers2 = {
                    'authorization': my_token,
                    'content-type': "application/json",
                    # 'content-type': "application/x-www-form-urlencoded",
                    'accept': "application/json"  # json
                }
                input_dict = {"lastName": "", "firstName": "", "address": "",
                              "city": "", "state": "", "zip": "", "referenceId": "ABC-xyz1",
                              "fields": ["name", "phone", "address", "email"]}
                if first_name:
                    input_dict['firstName'] = first_name
                if last_name:
                    input_dict['lastName'] = last_name
                if full_address:
                    input_dict['address'] = full_address
                text_data = ' '
                name = ''
                phone = ''
                email = ''
                address = ''

                phone_traced = []
                email_traced = []
                address_traced = []
                try:
                    json_body = json.dumps(input_dict)
                    response = requests.request("POST", url2, data=json_body, headers=headers2)
                    if response.status_code == 200:
                        response_got = response.json()
                        text_data = response_got['result']
                        if text_data:
                            if 'name' in text_data[0]:
                                name = text_data[0]['name'][0]['data']
                            if 'phone' in text_data[0]:
                                for p in text_data[0]['phone']:
                                    phone += p['type'] + ' : ' + p['number'] + '.'+'\n'
                                    phone_traced.append(p['number']+'.')
                                phone = phone[:-1]
                            if 'email' in text_data[0]:
                                for e in text_data[0]['email']:
                                    email += e['data'] + '.'+'\n'
                                    email_traced.append(e['data']+'.')
                                email = email[:-1]
                            if 'address' in text_data[0]:
                                for a in text_data[0]['address']:
                                    address += 'From: ' + a['dateRange'].split('-')[0] + ' To ' + \
                                               a['dateRange'].split('-')[1] + ' :\n' + a['data'] + '.'+'\n'
                                    address_traced.append('From: ' + a['dateRange'].split('-')[0] + ' To ' + \
                                               a['dateRange'].split('-')[1] + ' :\n' + a['data']+'.')
                                address = address[:-1]
                    single_obj.full_name = name
                    single_obj.phone = phone
                    single_obj.email = email
                    single_obj.address = address
                    single_obj.api_response = text_data
                    single_obj.is_validation_complete = True
                    single_obj.save()
                    for eml in email_traced:
                        EmailTraced.objects.create(skiptrace_id=single_obj.id, email=eml)
                    for phn in phone_traced:
                        PhoneTraced.objects.create(skiptrace_id=single_obj.id, phone=phn)
                    for addr in address_traced:
                        AddressTraced.objects.create(skiptrace_id=single_obj.id, address=addr)
                except:
                    single_obj.full_name = name
                    single_obj.phone = phone
                    single_obj.email = email
                    single_obj.address = address
                    single_obj.api_response = text_data
                    single_obj.is_validation_complete = True
                    single_obj.save()
                    for eml in email_traced:
                        EmailTraced.objects.create(skiptrace_id=single_obj.id, email=eml)
                    for phn in phone_traced:
                        PhoneTraced.objects.create(skiptrace_id=single_obj.id, phone=phn)
                    for addr in address_traced:
                        AddressTraced.objects.create(skiptrace_id=single_obj.id, address=addr)
                    print("in inner exception")
                    print(traceback.print_exc())

                document = Document()
                document.add_heading('SkipTraced Results', 0)
                document.add_heading('Full Name')
                document.add_paragraph(name)
                document.add_heading('Addresses Traced')
                document.add_paragraph(address)
                document.add_heading('Emails Traced')
                document.add_paragraph(email)
                document.add_heading('Phone Numbers Traced')
                document.add_paragraph(phone)
                # col_names = ('Full Name', 'Addresses Traced', 'Emails Traced', 'Phone Numbers Traced')
                # tbl = document.add_table(rows=1, cols=len(col_names))
                # hdr_cells = tbl.rows[0].cells
                # for idx, nm in enumerate(col_names):
                #     paragraph = hdr_cells[idx].paragraphs[0]
                #     run = paragraph.add_run(nm)
                #     run.bold = True
                # row_cells = tbl.add_row().cells
                # row_cells[0].text = name
                # row_cells[1].text = address
                # row_cells[2].text = email
                # row_cells[3].text = phone

                # document.add_page_break()
                base_path = os.path.join(BASE_DIR, 'media') + "/single_skiptrace"
                rand_num = str(random.sample(range(1000), 1)[0]) + "_" + str(random.sample(range(100), 1)[0])
                file_path = base_path + '/' + rand_num + 'single_skip_trace.docx'

                document.save(file_path)

                # fields = ['Full_Name', 'First_Name', 'Last_Name', 'Property_Address', 'Property_City', 'Property_State',
                #           'Property_Zip', 'Mailing_Address', 'Mailing_City', 'Mailing_State', 'Mailing_Zip',
                #           'Addresses_Traced', 'Emails_Traced', 'Phones_Traced']
                # header = dict((el, '') for el in fields)
                # header['Full_Name'] = [name]
                # header['First_Name'] = [first_name]
                # header['Last_Name'] = [last_name]
                # header['Property_Address'] = [full_address]
                # header['Property_City'] = [property_city]
                # header['Property_State'] = [property_state]
                # header['Property_Zip'] = [property_zip]
                # header['Mailing_Address'] = [mailing_address]
                # header['Mailing_City'] = [mailing_city]
                # header['Mailing_State'] = [mailing_state]
                # header['Mailing_Zip'] = [mailing_zip]
                # header['Addresses_Traced'] = [address]
                # header['Emails_Traced'] = [email]
                # header['Phones_Traced'] = [phone]

                # df = pd.DataFrame(header)
                # print("data frame we got: ", df)
                # writer = pd.ExcelWriter(file_path, engine='xlsxwriter')
                single_obj.file_name = rand_num + 'single_skip_trace.docx'
                single_obj.save()
                # df.to_excel(writer, sheet_name='Sheet1', index=False)
                # writer.save()
                if address != '':
                    print("success in skip tracing")  # address found
                    profile = UserProfile.objects.get(user_id=user)
                    prepaid_balance = PrepaidBalance.objects.get(user_id=user)
                    prepaid_balance.amount = prepaid_balance.amount - profile.skiptrace_price
                    prepaid_balance.save()
                    print("successfully deducted amount for skip trace")
                market_data = {
                    "name": name,
                    "first_name": first_name,
                    "last_name": last_name,
                    "address": address,
                    "full_address": full_address,
                    "email": email,
                    "phone": phone,
                    "file_name": '/media/single_skiptrace/'+rand_num + 'single_skip_trace.docx',
                }

                return Response({
                    "status": status.HTTP_200_OK,
                    "message": "Success",
                    "market_data": market_data
                })
            except:
                print(traceback.print_exc())
                return Response({
                    "status": status.HTTP_500_INTERNAL_SERVER_ERROR,
                    "message": "Error",
                    "market_data": []
                })


class ThirdPrtySkipTrace(APIView):
    # permission_classes = (IsAuthenticated,)
    def post(self, request):
        first_name = request.data.get('first_name')
        last_name = request.data.get('last_name')
        full_address = request.data.get('address')
        property_city = request.data.get('city')
        property_state = request.data.get('state')
        property_zip = request.data.get('zip')
        mailing_address = request.data.get('mailing_address')
        mailing_city = request.data.get('mailing_city')
        mailing_state = request.data.get('mailing_state')
        mailing_zip = request.data.get('mailing_zip')
        user = request.data.get('user')
        print("Third Party data is: ", first_name, last_name, full_address, property_city, property_state, property_zip, user, mailing_address, mailing_city, mailing_state, mailing_zip)
        if full_address:
            try:
                url1 = "https://login-api.idicore.com/apiclient"
                payload = "{\"glba\":\"otheruse\",\"dppa\":\"none\"}"
                d = base64.b64encode(
                    b'api-client@acesolutions:ELw%7UYUbfkb5V!naYXhksX2Tu3LYvA%dqvA5YkevtXU%4x59GLnvhnJ5kqiL4Bi')
                print("d is: ", d)
                abc = "Basic " + str(d.decode("utf-8"))
                print("abc is: ", abc)
                headers = {
                    'authorization': abc,
                    'content-type': "application/json"
                }
                response = requests.request("POST", url1, data=payload, headers=headers)
                my_token = response.text  # token created for 15 minutes
                print("skkiptrace token: ", my_token)
                url2 = "https://api.idicore.com/search"
                headers2 = {
                    'authorization': my_token,
                    'content-type': "application/json",
                    # 'content-type': "application/x-www-form-urlencoded",
                    'accept': "application/json"  # json
                }
                input_dict = {"lastName": "", "firstName": "", "address": "",
                              "city": "", "state": "", "zip": "", "referenceId": "ABC-xyz1",
                              "fields": ["name", "phone", "address", "email"]}
                if first_name:
                    input_dict['firstName'] = first_name
                if last_name:
                    input_dict['lastName'] = last_name
                if full_address:
                    input_dict['address'] = full_address
                text_data = ' '
                name = ''
                phone = ''
                email = ''
                address = ''

                phone_traced = []
                email_traced = []
                address_traced = []
                try:
                    json_body = json.dumps(input_dict)
                    response = requests.request("POST", url2, data=json_body, headers=headers2)
                    if response.status_code == 200:
                        response_got = response.json()
                        text_data = response_got['result']
                        if text_data:
                            if 'name' in text_data[0]:
                                name = text_data[0]['name'][0]['data']
                            if 'phone' in text_data[0]:
                                for p in text_data[0]['phone']:
                                    phone += p['type'] + ' : ' + p['number'] + '.'+'\n'
                                    phone_traced.append(p['number']+'.')
                                phone = phone[:-1]
                            if 'email' in text_data[0]:
                                for e in text_data[0]['email']:
                                    email += e['data'] + '.'+'\n'
                                    email_traced.append(e['data']+'.')
                                email = email[:-1]
                            if 'address' in text_data[0]:
                                for a in text_data[0]['address']:
                                    address += 'From: ' + a['dateRange'].split('-')[0] + ' To ' + \
                                               a['dateRange'].split('-')[1] + ' :\n' + a['data'] + '.'+'\n'
                                    address_traced.append('From: ' + a['dateRange'].split('-')[0] + ' To ' + \
                                               a['dateRange'].split('-')[1] + ' :\n' + a['data']+'.')
                                address = address[:-1]

                except:
                    print("in inner exception")
                    print(traceback.print_exc())

                market_data = {
                    "name": name,
                    "first_name": first_name,
                    "last_name": last_name,
                    "address": address,
                    "full_address": full_address,
                    "email": email,
                    "phone": phone,
                    "phone_traced": phone_traced,
                    "email_traced": email_traced,
                    "address_traced": address_traced,
                    "text_data": text_data,
                }

                return Response({
                    "status": status.HTTP_200_OK,
                    "message": "Success",
                    "market_data": market_data
                })
            except:
                print(traceback.print_exc())
                return Response({
                    "status": status.HTTP_500_INTERNAL_SERVER_ERROR,
                    "message": "Error",
                    "market_data": []
                })


@user_has_single_skip_trace_Balance
def single_skip_trace_view(request):
    return render(request, 'skiptrace/single_skiptrace.html', {"user": request.user.id,
                                                               })


def skip_trace_existing_file_view(request):
    list_data = []
    lists = File.objects.filter(user=request.user)
    for lis in lists:
        data = {
            "id": lis.id,
            "file_name": lis.file_name,
            "list_name": lis.list.list_name+' ('+lis.file_name[7:]+')',
        }
        list_data.append(data)
    return render(request, 'skiptrace/skiptrace_existing_list.html', {"user": request.user.id,
                                                                      "list_data": list_data})


def upload_skip_trace_file(request):
    file = request.FILES.get("file")
    if file.name.endswith('.csv'):
        try:
            decoded_file = file.read().decode('utf-8').splitlines()
            reader = csv.DictReader(decoded_file)
            row_count = (len(list(reader)) - 1)
        except:
            print(traceback.print_exc())
            temp = {
                "sheet_name": "",
                "header": "",
                "row_count": False,
                "total_price": "",
                "balance": "",
                "acc_balance": "",
                "file_name": ""
            }

            return JsonResponse(temp)
        if row_count > 49:
            total_price = row_count * 0.15
            pass

        else:
            # raise LookupError('Please upload a list having at least 50 records in it')
            temp = {
                "sheet_name": "",
                "header": "",
                "row_count": row_count,
                "total_price": "",
                "balance": "",
                "acc_balance": "",
                "file_name": ""
            }

            return JsonResponse(temp)

    else:
        try:
            wb = openpyxl.load_workbook(file, enumerate)
            sheet = wb.worksheets[0]
            row_count = (sheet.max_row - 1)
        except:
            print(traceback.print_exc())
            temp = {
                "sheet_name": "",
                "header": "",
                "row_count": False,
                "total_price": "",
                "balance": "",
                "acc_balance": "",
                "file_name": ""
            }

            return JsonResponse(temp)
        if row_count > 49:
            total_price = row_count * 0.15

        else:
            temp = {
                "sheet_name": "",
                "header": "",
                "row_count": row_count,
                "total_price": "",
                "balance": "",
                "acc_balance": "",
                "file_name": ""
            }

            return JsonResponse(temp)
            # raise LookupError('Please upload a list having at least 50 records in it')

    try:
        path = settings.MEDIA_ROOT + '/skiptrace_files'

        fs = FileSystemStorage(path, None, file_permissions_mode=0o644)
        id_ = str(random.sample(range(1000), 1)[0]) + "_" + str(random.sample(range(100), 1)[0])
        filename = id_ + "_" + file.name
        fs.save(filename, file)
        old_file = filename
        path = os.path.join(BASE_DIR, 'media') + "/skiptrace_files"
        file_path = path + '/' + filename

        ar = []
        sheets = []
        if ".csv" in file.name:
            df = pd.read_csv(file_path)
            name_old = os.path.splitext(filename)[0]
            filename = name_old+'.xlsx'
            new_name = path + '/' + filename
            df.to_excel(new_name, index=None, header=True)
            sheets.append("WorkSheet")
            file_path = new_name
            os.remove(os.path.join(BASE_DIR, 'media/skiptrace_files/') +old_file)
        xl = pd.ExcelFile(file_path)
        sheet_name = xl.sheet_names
        for sheet in sheet_name:
            sheets.append(sheet)

        for i in pd.read_excel(file_path):
            if "Unnamed:" not in i:
                ar.append(i)
        balance = PrepaidBalance.objects.get(user=request.user).amount
        acc_balance = balance >= total_price
        temp = {
            "sheet_name": sheets,
            "header": ar,
            "row_count": row_count,
            "total_price": round(decimal.Decimal(total_price), 3),
            "balance": balance,
            "acc_balance": acc_balance,
            "file_name": filename
        }

        return JsonResponse(temp)
    except Exception as e:
        print(e)
        print(traceback.print_exc())
        raise Exception('not save in save image in folder!')


def get_col_by_skip_trace_sheet_name(request):
    path = os.path.join(BASE_DIR, 'media') + "/skiptrace_files"
    file_path = path + '/' + request.POST["file_name"]
    ar = []
    sheets = []
    if ".csv" in request.POST["file_name"]:
        df = pd.read_csv(file_path)
        for data in df:
            ar.append(data)
    else:
        for i in pd.read_excel(file_path, sheet_name=request.POST["sheet_name"]):
            if "Unnamed:" not in i:
                ar.append(i)
    temp = {

        "header": ar,
        "file_name": request.POST["file_name"]
    }
    return JsonResponse(temp)


def add_new_skip_trace_list(request):
    user = request.user
    sheet_name = request.POST["sheet_name"]
    file_name = request.POST["file_name"]
    source_fields = request.POST["source_field"]

    if 'firstname' in source_fields and 'lastname' in source_fields and 'propertyaddress' in source_fields and 'propertycity' in source_fields:
        if 'propertystate' in source_fields or 'propertyzip' in source_fields:
            try:
                fields_list = source_fields.split(",")
                fields = ''
                for f in fields_list:
                    if fields == '':
                        fields += f
                    else:
                        fields += ','+f
                with transaction.atomic():
                    skip_trace = SkipTraceFile.objects.create(user=user, file_name=file_name,
                                                              sheet_name="sheet1", destination_fields=fields)
                    temp = {
                        "status": 200,
                        "message": "List Submitted Successfully"
                    }
                    return JsonResponse(temp)
            except:
                print(traceback.print_exc())
                temp = {
                    "status": 500,
                }
                return JsonResponse(temp)
        else:
            temp = {
                "status": 201,
            }
            return JsonResponse(temp)
    else:
        temp = {
            "status": 201,
        }
        return JsonResponse(temp)


class GetHitsCount(APIView):
    # permission_classes = (IsAuthenticated,)
    def post(self, request):
        user = request.data.get('user')
        skip_trace = SkipTraceFile.objects.filter(user=user)
        skip_data = []
        if skip_trace:
            try:
                for skip in skip_trace:
                    data = {
                        "id": 'percent'+str(skip.id),
                        "ids": skip.id,
                        "hits_percentage": skip.hits_percentage,
                        "total_records": skip.total_records,
                        "total_hits": skip.total_hits,
                        "exported_link": '/media/exported_skiptrace/' + skip.file_name if skip.status == 'Uploaded' else 'javascript:void(0)',
                        "status": skip.status
                    }
                    skip_data.append(data)
                return Response({"status": status.HTTP_200_OK, "data": skip_data})
            except:
                print(traceback.print_exc())
                return Response({"status": status.HTTP_500_INTERNAL_SERVER_ERROR, "message": "Server error"})
        else:
            return Response({"status": status.HTTP_200_OK, "data": skip_data})


class AddExistingFileForSkipTrace(APIView):
    # permission_classes = (IsAuthenticated,)
    def post(self, request):
        file_name = request.data.get('list_id')
        user = request.data.get('user')
        balance = PrepaidBalance.objects.get(user=user).amount
        from_file_path = os.path.join(BASE_DIR, 'media') + "/files/"+ file_name
        to_path = os.path.join(BASE_DIR, 'media') + "/skiptrace_files"
        if file_name:
            try:
                if file_name.endswith('.csv'):
                    decoded_file = from_file_path.read().decode('utf-8').splitlines()
                    reader = csv.DictReader(decoded_file)
                    row_count = (len(list(reader))-1)
                    total_price = row_count * 0.15
                    acc_balance = balance >= total_price
                    if acc_balance:
                        if row_count > 49:
                            pass
                        else:
                            return Response({"status": status.HTTP_404_NOT_FOUND,
                                             "message": "Please upload a list having at least 50 records in it"})
                    else:
                        return Response({"status": status.HTTP_404_NOT_FOUND,
                                         "message": "You don't have sufficient balance to Skip Trace this File"})
                else:
                    excel_workbook = openpyxl.load_workbook(from_file_path, read_only=True)
                    sheet = excel_workbook.worksheets[0]
                    row_count = (sheet.max_row - 1)
                    total_price = row_count * 0.15
                    acc_balance = balance >= total_price
                    if acc_balance:
                        if row_count > 49:
                            pass
                        else:
                            return Response({"status": status.HTTP_404_NOT_FOUND,
                                             "message": "Please upload a list having at least 50 records in it"})
                    else:
                        return Response({"status": status.HTTP_404_NOT_FOUND,
                                         "message": "You don't have sufficient balance to Skip Trace this File"})

                shutil.copy(from_file_path, to_path)
                media_file = to_path+'/'+file_name
                ar = []
                sheets = []
                if ".csv" in file_name:
                    df = pd.read_csv(media_file)
                    sheets.append("WorkSheet")
                    for data in df:
                        ar.append(data)
                else:
                    xl = pd.ExcelFile(media_file)
                    sheet_name = xl.sheet_names
                    for sheet in sheet_name:
                        sheets.append(sheet)
                    print("==================", xl.sheet_names)

                    for i in pd.read_excel(media_file):
                        if "Unnamed:" not in i:
                            ar.append(i)
                balance = PrepaidBalance.objects.get(user=user).amount
                acc_balance = balance >= total_price
                if acc_balance:
                    response = {
                        "sheet_name": sheets,
                        "header": ar,
                        "row_count": row_count,
                        "total_price": round(decimal.Decimal(total_price), 3),
                        "balance": balance,
                        "acc_balance": acc_balance,
                        "file_name": file_name
                    }
                    print("returning without going in main function: ", response)
                    return Response({"status": status.HTTP_200_OK, "response": response})
                else:
                    return Response({"status": status.HTTP_404_NOT_FOUND,
                                     "message": "You don't have sufficient balance to Skip Trace this File"})
            except:
                print(traceback.format_exc())
                return Response({"status": status.HTTP_404_NOT_FOUND,
                                 "message": "File is not ready to be used or something is wrong with the file!"})


class Add_New_Payment_View(APIView):
    def post(self, request):
        balance = request.POST.get('balance')
        user = request.POST.get('user')
        key = settings.STRIPE_PUBLISHABLE_KEY
        if balance and user:
            email = User.objects.get(id=user).email
            content = {'key': key, 'balance': float(balance) * 100.0, 'email': email}
            return Response({"status": status.HTTP_200_OK,
                             "response": content})
        else:
            return Response({"status": status.HTTP_400_BAD_REQUEST,
                             "response": ""})
